import io
import re
from functools import lru_cache
from pathlib import Path

from docx import Document
from pypdf import PdfReader

from asem_talent.domain.models import ResumeContext, ResumeInternship, ResumeParseResponse, ResumeProject

MAX_UPLOAD_BYTES = 2_000_000
MAX_EXTRACTED_CHARS = 12_000
MAX_PREVIEW_CHARS = 1_000
MIN_DIRECT_PDF_TEXT_CHARS = 24
OCR_RENDER_SCALE = 2.0

SUPPORTED_EXTENSIONS = {".pdf", ".docx"}

EMAIL_PATTERN = re.compile(r"\b[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[A-Za-z]{2,}\b")
PHONE_PATTERN = re.compile(r"(?<!\d)(?:\+?\d[\d\s().-]{7,}\d)")
URL_PATTERN = re.compile(r"https?://\S+|www\.\S+")
NRIC_PATTERN = re.compile(r"\b\d{6}-\d{2}-\d{4}\b")

SKILL_KEYWORDS = {
    "python_basics": ["python", "pandas", "numpy", "jupyter"],
    "c++": ["c++", "cpp"],
    "debugging": ["debugging", "troubleshooting", "fault isolation", "root cause analysis"],
    "data_acquisition": ["data acquisition", "daq", "data logging", "telemetry"],
    "embedded_systems": ["embedded", "firmware", "microcontroller", "arduino", "stm32"],
    "sensor_integration": ["sensor integration", "sensors", "vision system", "camera calibration"],
    "test_automation": ["test automation", "automated testing", "test script", "pytest"],
    "robotics_integration": ["robotics", "robot arm", "robot integration", "machine vision"],
    "process_discipline": ["spc", "process control", "process improvement", "yield"],
    "plc_programming": ["plc", "ladder logic", "siemens s7"],
}

TOOL_KEYWORDS = {
    "oscilloscope": ["oscilloscope", "scope"],
    "labview": ["labview"],
    "matlab": ["matlab", "simulink"],
    "autocad": ["autocad", "solidworks", "cad"],
    "excel": ["excel"],
}

ROLE_SIGNAL_KEYWORDS = {
    "validation engineer trainee": ["validation", "wafer test", "test engineer"],
    "robotics trainee": ["robotics", "automation cell", "robot arm"],
    "embedded systems ojt": ["embedded", "firmware", "microcontroller"],
    "process automation trainee": ["process control", "yield", "spc"],
}

CERTIFICATION_KEYWORDS = {
    "ipc_basics": ["ipc", "ipc-a-610"],
    "lean_six_sigma": ["six sigma", "green belt", "lean"],
    "labview_certification": ["labview certification"],
}

SECTION_HEADINGS = {
    "projects": {"projects", "project experience", "academic projects", "selected projects"},
    "experience": {"experience", "work experience", "internship", "internships", "employment"},
    "certifications": {"certifications", "certificates", "licenses"},
}


class ResumeParseError(ValueError):
    def __init__(self, detail: str, status_code: int = 400) -> None:
        super().__init__(detail)
        self.detail = detail
        self.status_code = status_code


def parse_resume_document(file_name: str, content_type: str | None, content_bytes: bytes) -> ResumeParseResponse:
    extension = Path(file_name).suffix.lower()
    if extension not in SUPPORTED_EXTENSIONS:
        raise ResumeParseError("Only PDF and DOCX resume uploads are supported.", status_code=415)
    if not content_bytes:
        raise ResumeParseError("Uploaded file is empty.", status_code=400)
    if len(content_bytes) > MAX_UPLOAD_BYTES:
        raise ResumeParseError(
            f"Resume upload exceeds the {MAX_UPLOAD_BYTES} byte limit.",
            status_code=413,
        )

    raw_text, ocr_used = _extract_text(extension, content_bytes)
    if not raw_text.strip():
        raise ResumeParseError("No extractable text was found in the uploaded resume.", status_code=400)

    normalized_text = _normalize_text(raw_text)
    truncated = len(normalized_text) > MAX_EXTRACTED_CHARS
    if truncated:
        normalized_text = normalized_text[:MAX_EXTRACTED_CHARS]

    redacted_text, redaction_counts = _redact_pii(normalized_text)
    resume_context = _build_resume_context(redacted_text)
    warnings: list[str] = []
    if truncated:
        warnings.append("Resume text exceeded the parser limit and was truncated.")
    if ocr_used:
        warnings.append("OCR fallback was used because the PDF did not expose enough extractable text.")
    if not resume_context.skill_tags and not resume_context.project_highlights and not resume_context.internship_highlights:
        warnings.append("The parser extracted limited structured evidence. Review the resume_context before using it in a live decision.")

    return ResumeParseResponse(
        file_name=file_name,
        file_type=extension.lstrip("."),
        file_size_bytes=len(content_bytes),
        redacted_preview=redacted_text[:MAX_PREVIEW_CHARS],
        redaction_counts=redaction_counts,
        truncated=truncated,
        warnings=warnings,
        resume_context=resume_context,
    )


def _extract_text(extension: str, content_bytes: bytes) -> tuple[str, bool]:
    if extension == ".pdf":
        direct_text = _extract_pdf_text(content_bytes)
        if _should_use_pdf_ocr(direct_text):
            ocr_text = _extract_pdf_text_with_ocr(content_bytes)
            if ocr_text.strip():
                return ocr_text, True
        return direct_text, False
    if extension == ".docx":
        return _extract_docx_text(content_bytes), False
    raise ResumeParseError("Unsupported resume file type.", status_code=415)


def _extract_pdf_text(content_bytes: bytes) -> str:
    try:
        reader = PdfReader(io.BytesIO(content_bytes))
    except Exception as error:  # pragma: no cover - library-specific failures
        raise ResumeParseError(f"PDF parsing failed: {error}", status_code=400) from error

    text_parts = [page.extract_text() or "" for page in reader.pages]
    return "\n".join(text_parts)


def _should_use_pdf_ocr(text: str) -> bool:
    alphanumeric_chars = sum(character.isalnum() for character in text)
    return alphanumeric_chars < MIN_DIRECT_PDF_TEXT_CHARS


@lru_cache(maxsize=1)
def _get_ocr_engine():
    try:
        from rapidocr_onnxruntime import RapidOCR
    except ImportError as error:  # pragma: no cover - dependency issue
        raise ResumeParseError(
            "OCR support requires rapidocr-onnxruntime to be installed in the runtime environment.",
            status_code=500,
        ) from error
    return RapidOCR()


def _extract_pdf_text_with_ocr(content_bytes: bytes) -> str:
    try:
        import pypdfium2 as pdfium
    except ImportError as error:  # pragma: no cover - dependency issue
        raise ResumeParseError(
            "OCR support requires pypdfium2 to be installed in the runtime environment.",
            status_code=500,
        ) from error

    try:
        document = pdfium.PdfDocument(io.BytesIO(content_bytes))
    except Exception as error:  # pragma: no cover - library-specific failures
        raise ResumeParseError(f"PDF OCR preparation failed: {error}", status_code=400) from error

    ocr_engine = _get_ocr_engine()
    extracted_pages: list[str] = []
    try:
        page_total = len(document)
        for page_index in range(page_total):
            page = document.get_page(page_index)
            bitmap = page.render(scale=OCR_RENDER_SCALE)
            image = bitmap.to_numpy()
            if len(getattr(image, "shape", ())) == 3 and image.shape[2] > 3:
                image = image[:, :, :3]
            ocr_result, _ = ocr_engine(image)
            page_text = _flatten_ocr_result(ocr_result)
            if page_text:
                extracted_pages.append(page_text)
            page.close()
    except ResumeParseError:
        raise
    except Exception as error:  # pragma: no cover - library-specific failures
        raise ResumeParseError(f"PDF OCR failed: {error}", status_code=400) from error
    finally:
        document.close()

    return "\n".join(extracted_pages)


def _flatten_ocr_result(ocr_result) -> str:
    if not ocr_result:
        return ""
    text_parts = []
    for item in ocr_result:
        if isinstance(item, (list, tuple)) and len(item) >= 2 and isinstance(item[1], str):
            cleaned = item[1].strip()
            if cleaned:
                text_parts.append(cleaned)
    return "\n".join(text_parts)


def _extract_docx_text(content_bytes: bytes) -> str:
    try:
        document = Document(io.BytesIO(content_bytes))
    except Exception as error:  # pragma: no cover - library-specific failures
        raise ResumeParseError(f"DOCX parsing failed: {error}", status_code=400) from error

    paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    table_cells = []
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    table_cells.append(cell.text)
    return "\n".join(paragraphs + table_cells)


def _normalize_text(text: str) -> str:
    cleaned = text.replace("\r", "\n")
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)
    cleaned = re.sub(r"[ \t]{2,}", " ", cleaned)
    return cleaned.strip()


def _redact_pii(text: str) -> tuple[str, dict[str, int]]:
    counts = {"emails": 0, "phones": 0, "urls": 0, "nric": 0}

    def replace_with_count(pattern: re.Pattern[str], replacement: str, key: str, value: str) -> str:
        matches = pattern.findall(value)
        counts[key] += len(matches)
        return pattern.sub(replacement, value)

    redacted = text
    redacted = replace_with_count(EMAIL_PATTERN, "[REDACTED_EMAIL]", "emails", redacted)
    redacted = replace_with_count(PHONE_PATTERN, "[REDACTED_PHONE]", "phones", redacted)
    redacted = replace_with_count(URL_PATTERN, "[REDACTED_URL]", "urls", redacted)
    redacted = replace_with_count(NRIC_PATTERN, "[REDACTED_ID]", "nric", redacted)
    return redacted, counts


def _build_resume_context(text: str) -> ResumeContext:
    lines = [line.strip(" -•\t") for line in text.splitlines() if line.strip()]
    normalized_text = text.lower()
    sections = _extract_sections(lines)

    project_lines = sections["projects"] or _first_matching_lines(lines, {"project", "prototype", "capstone"}, limit=3)
    experience_lines = sections["experience"] or _first_matching_lines(lines, {"intern", "engineer", "assistant", "technician"}, limit=2)
    certification_lines = sections["certifications"] or _first_matching_lines(lines, {"cert", "sigma", "ipc"}, limit=4)

    summary_source = []
    for line in lines:
        lowercase_line = line.lower()
        if "@" in line or "[redacted_" in lowercase_line:
            continue
        if any(token in lowercase_line for token in ("project", "experience", "certification", "skills")):
            continue
        summary_source.append(line)
        if len(" ".join(summary_source)) >= 320:
            break

    projects = [_parse_project_line(line) for line in project_lines[:3]]
    internships = [_parse_internship_line(line) for line in experience_lines[:2]]
    projects = [project for project in projects if project.title]
    internships = [internship for internship in internships if internship.role_title or internship.organization]

    return ResumeContext(
        summary=" ".join(summary_source)[:400] or None,
        skill_tags=_find_tags(normalized_text, SKILL_KEYWORDS),
        tool_tags=_find_tags(normalized_text, TOOL_KEYWORDS),
        project_highlights=projects,
        internship_highlights=internships,
        certifications=_extract_certifications(certification_lines, normalized_text),
        inferred_role_signals=_find_tags(normalized_text, ROLE_SIGNAL_KEYWORDS),
    )


def _extract_sections(lines: list[str]) -> dict[str, list[str]]:
    sections = {"projects": [], "experience": [], "certifications": []}
    current_section: str | None = None
    for line in lines:
        normalized = line.lower().rstrip(":")
        heading = _match_heading(normalized)
        if heading is not None:
            current_section = heading
            continue
        if current_section is not None:
            sections[current_section].append(line)
    return sections


def _match_heading(normalized_line: str) -> str | None:
    for section_name, headings in SECTION_HEADINGS.items():
        if normalized_line in headings:
            return section_name
    return None


def _first_matching_lines(lines: list[str], keywords: set[str], limit: int) -> list[str]:
    matches = []
    for line in lines:
        lower_line = line.lower()
        if any(keyword in lower_line for keyword in keywords):
            matches.append(line)
        if len(matches) >= limit:
            break
    return matches


def _find_tags(text: str, mapping: dict[str, list[str]]) -> list[str]:
    found: list[str] = []
    for tag, keywords in mapping.items():
        if any(keyword in text for keyword in keywords):
            found.append(tag)
    return found


def _parse_project_line(line: str) -> ResumeProject:
    title, summary = _split_title_and_summary(line)
    line_text = f"{title} {summary or ''}".lower()
    return ResumeProject(
        title=title,
        summary=summary,
        skill_tags=_find_tags(line_text, SKILL_KEYWORDS),
        outcome_tags=_extract_outcome_tags(line_text),
    )


def _parse_internship_line(line: str) -> ResumeInternship:
    role_title = line
    organization = "UNSPECIFIED"
    lowered = line.lower()
    if " at " in lowered:
        split_index = lowered.index(" at ")
        role_title = line[:split_index].strip(" -|")
        organization = line[split_index + 4 :].strip(" -|")
    elif " - " in line:
        first, second = line.split(" - ", 1)
        if any(token in first.lower() for token in ("intern", "engineer", "assistant", "technician")):
            role_title = first.strip()
            organization = second.strip()
        else:
            organization = first.strip()
            role_title = second.strip()
    return ResumeInternship(
        organization=organization,
        role_title=role_title,
        summary=None,
        skill_tags=_find_tags(line.lower(), SKILL_KEYWORDS),
    )


def _split_title_and_summary(line: str) -> tuple[str, str | None]:
    for separator in (":", " - ", " | "):
        if separator in line:
            title, summary = line.split(separator, 1)
            return title.strip(), summary.strip() or None
    return line.strip(), None


def _extract_outcome_tags(text: str) -> list[str]:
    outcomes = []
    if "yield" in text:
        outcomes.append("yield_improvement")
    if "defect" in text:
        outcomes.append("defect_detection")
    if "automation" in text:
        outcomes.append("automation_delivery")
    return outcomes


def _extract_certifications(certification_lines: list[str], normalized_text: str) -> list[str]:
    found = _find_tags(normalized_text, CERTIFICATION_KEYWORDS)
    for line in certification_lines:
        cleaned = line.strip()
        if cleaned and cleaned.lower() not in {item.lower() for item in found}:
            found.append(cleaned[:80])
    return found[:5]